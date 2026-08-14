"""Bounded, local text extraction for supported document formats."""

from __future__ import annotations

import io
import os
import zipfile
from collections.abc import Mapping
from dataclasses import replace
from pathlib import Path, PurePosixPath

from docx import Document
from pypdf import PdfReader

from desktop_agent.models import (
    ContentExtractionStatus,
    FileMetadata,
)
from desktop_agent.entry_snapshot import SnapshotLimitError, snapshot_directory

MAX_CONTENT_CHARS = 4_000
MAX_FILE_BYTES = 10 * 1024 * 1024
MAX_ARCHIVE_BYTES = 512 * 1024 * 1024
MAX_PDF_PAGES = 10
MAX_CONTAINER_DEPTH = 3
MAX_CONTAINER_ENTRIES = 30
MAX_CONTAINER_SCAN_ENTRIES = 50_000
MAX_ARCHIVE_MEMBERS = 100_000
MAX_CONTAINER_UNCOMPRESSED_BYTES = 20 * 1024 * 1024
MAX_CONTAINER_DOCUMENTS = 6

IGNORED_CONTAINER_DIRECTORIES = {
    ".git",
    ".idea",
    ".mypy_cache",
    ".pytest_cache",
    ".ruff_cache",
    ".venv",
    ".vscode",
    "__macosx",
    "__pycache__",
    "build",
    "dist",
    "node_modules",
    "site-packages",
    "venv",
}
ACADEMIC_FILENAME_MARKERS = (
    "实验报告",
    "课程设计",
    "大作业",
    "课程报告",
    "report",
)

PLAIN_TEXT_EXTENSIONS = {
    ".csv",
    ".html",
    ".json",
    ".log",
    ".md",
    ".py",
    ".rst",
    ".text",
    ".tsv",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
}
SUPPORTED_EXTENSIONS = PLAIN_TEXT_EXTENSIONS | {".docx", ".pdf", ".zip"}


def extract_contents(
    files_by_id: Mapping[str, FileMetadata],
) -> dict[str, FileMetadata]:
    """Return new metadata snapshots enriched with bounded text excerpts."""

    return {
        file_id: extract_content(file)
        for file_id, file in files_by_id.items()
    }


def extract_content(file: FileMetadata) -> FileMetadata:
    """Extract a safe, bounded excerpt without changing the source file."""

    if not file.is_directory and file.extension not in SUPPORTED_EXTENSIONS:
        return replace(
            file,
            content_status=ContentExtractionStatus.UNSUPPORTED,
            content_excerpt=None,
        )

    size_limit = (
        MAX_ARCHIVE_BYTES if file.extension == ".zip" else MAX_FILE_BYTES
    )
    if not file.is_directory and file.size_bytes > size_limit:
        return replace(
            file,
            content_status=ContentExtractionStatus.TOO_LARGE,
            content_excerpt=None,
        )

    if not _source_is_unchanged(file):
        return replace(
            file,
            content_status=ContentExtractionStatus.ERROR,
            content_excerpt=None,
        )

    try:
        if file.is_directory:
            text = _read_directory(file.path)
        elif file.extension in PLAIN_TEXT_EXTENSIONS:
            text = _read_plain_text(file)
        elif file.extension == ".pdf":
            text = _read_pdf(file)
        elif file.extension == ".docx":
            text = _read_docx(file)
        else:
            text = _read_zip(file.path)
    except Exception:
        # Third-party parsers expose many format-specific exceptions. A bad
        # document must degrade to metadata-only classification, not stop the
        # whole organization workflow.
        return replace(
            file,
            content_status=ContentExtractionStatus.ERROR,
            content_excerpt=None,
        )

    return replace(
        file,
        content_status=ContentExtractionStatus.EXTRACTED,
        content_excerpt=_bounded_excerpt(text),
        academic_material_hint=(
            "实验报告"
            if "实验报告" in f"{file.name}\n{text}"
            else file.academic_material_hint
        ),
    )


def _source_is_unchanged(file: FileMetadata) -> bool:
    """Refuse symlinks and files changed since the metadata scan."""

    path = file.path
    if path.is_symlink() or not (path.is_file() or path.is_dir()):
        return False
    if file.is_directory:
        if not path.is_dir() or file.tree_fingerprint is None:
            return False
        try:
            size_bytes, fingerprint = snapshot_directory(path)
        except (OSError, SnapshotLimitError):
            return False
        return (
            size_bytes == file.size_bytes
            and fingerprint == file.tree_fingerprint
        )
    if not path.is_file():
        return False
    try:
        file_stat = path.stat(follow_symlinks=False)
    except OSError:
        return False
    if file_stat.st_size != file.size_bytes:
        return False
    return abs(file_stat.st_mtime - file.modified_at.timestamp()) <= 0.001


def _read_plain_text(file: FileMetadata) -> str:
    """Read only enough text to build the configured excerpt."""

    with file.path.open("r", encoding="utf-8", errors="replace") as stream:
        return stream.read(MAX_CONTENT_CHARS + 1)


def _read_pdf(file: FileMetadata) -> str:
    """Extract a bounded number of PDF pages locally."""

    reader = PdfReader(file.path)
    parts: list[str] = []
    current_length = 0

    for page in reader.pages[:MAX_PDF_PAGES]:
        page_text = page.extract_text() or ""
        parts.append(page_text)
        current_length += len(page_text)
        if current_length >= MAX_CONTENT_CHARS:
            break

    return "\n".join(parts)


def _read_docx(file: FileMetadata) -> str:
    """Extract paragraph text from a Word document locally."""

    document = Document(file.path)
    parts: list[str] = []
    current_length = 0

    for paragraph in document.paragraphs:
        if not paragraph.text:
            continue
        parts.append(paragraph.text)
        current_length += len(paragraph.text)
        if current_length >= MAX_CONTENT_CHARS:
            break

    return "\n".join(parts)


def _read_directory(path: Path) -> str:
    """Inspect a directory manifest and selected documents without recursion drift."""

    candidates: list[Path] = []
    scanned_entries = 0
    for current_root, directory_names, file_names in os.walk(
        path,
        topdown=True,
        followlinks=False,
    ):
        current = Path(current_root)
        relative_depth = len(current.relative_to(path).parts)
        directory_names[:] = sorted(
            name
            for name in directory_names
            if not name.startswith(".")
            and name.casefold() not in IGNORED_CONTAINER_DIRECTORIES
            and not (current / name).is_symlink()
            and relative_depth < MAX_CONTAINER_DEPTH
        )
        for name in sorted(file_names):
            scanned_entries += 1
            if scanned_entries > MAX_CONTAINER_SCAN_ENTRIES:
                break
            entry = current / name
            if (
                name.startswith(".")
                or name.startswith("~$")
                or entry.is_symlink()
                or not entry.is_file()
            ):
                continue
            if len(entry.relative_to(path).parts) > MAX_CONTAINER_DEPTH:
                continue
            candidates.append(entry)
        if scanned_entries > MAX_CONTAINER_SCAN_ENTRIES:
            break

    candidates.sort(
        key=lambda entry: _candidate_priority(
            entry.relative_to(path).as_posix(), entry.suffix.lower()
        )
    )
    entries = candidates[:MAX_CONTAINER_ENTRIES]

    parts = ["容器文件清单："]
    parts.extend(f"- {entry.relative_to(path).as_posix()}" for entry in entries)
    extracted_documents = 0
    inspected_bytes = 0
    for entry in entries:
        extension = entry.suffix.lower()
        if extension not in (PLAIN_TEXT_EXTENSIONS | {".docx", ".pdf"}):
            continue
        size = entry.stat(follow_symlinks=False).st_size
        if size > MAX_FILE_BYTES:
            continue
        inspected_bytes += size
        if inspected_bytes > MAX_CONTAINER_UNCOMPRESSED_BYTES:
            break
        try:
            text = _read_document_path(entry, extension)
        except Exception:
            continue
        if text.strip():
            parts.append(
                f"\n[{entry.relative_to(path).as_posix()}]\n{text.strip()}"
            )
            extracted_documents += 1
        if (
            extracted_documents >= MAX_CONTAINER_DOCUMENTS
            or len("\n".join(parts)) >= MAX_CONTENT_CHARS
        ):
            break
    return "\n".join(parts)


def _read_zip(path: Path) -> str:
    """Inspect a ZIP in memory with traversal and decompression limits."""

    parts = ["压缩包文件清单："]
    extracted_documents = 0
    total_uncompressed = 0
    with zipfile.ZipFile(path) as archive:
        members = [member for member in archive.infolist() if not member.is_dir()]
        if len(members) > MAX_ARCHIVE_MEMBERS:
            raise ValueError("archive contains too many entries")

        safe_members: list[tuple[zipfile.ZipInfo, PurePosixPath]] = []
        for member in members:
            member_path = PurePosixPath(_display_zip_name(member.filename))
            if (
                member_path.is_absolute()
                or ".." in member_path.parts
                or len(member_path.parts) > MAX_CONTAINER_DEPTH
                or any(part.startswith(".") for part in member_path.parts)
                or any(
                    part.casefold() in IGNORED_CONTAINER_DIRECTORIES
                    for part in member_path.parts[:-1]
                )
                or member_path.name.startswith("~$")
            ):
                continue
            if member.flag_bits & 0x1:
                continue
            safe_members.append((member, member_path))

        safe_members.sort(
            key=lambda item: _candidate_priority(
                item[1].as_posix(), Path(item[1].name).suffix.lower()
            )
        )
        for member, member_path in safe_members[:MAX_CONTAINER_ENTRIES]:
            parts.append(f"- {member_path.as_posix()}")
            extension = Path(member_path.name).suffix.lower()
            if extension not in (PLAIN_TEXT_EXTENSIONS | {".docx", ".pdf"}):
                continue
            if member.file_size > MAX_FILE_BYTES:
                continue
            if (
                total_uncompressed + member.file_size
                > MAX_CONTAINER_UNCOMPRESSED_BYTES
            ):
                continue
            with archive.open(member) as stream:
                data = stream.read(MAX_FILE_BYTES + 1)
            if len(data) > MAX_FILE_BYTES:
                continue
            total_uncompressed += len(data)
            try:
                text = _read_document_bytes(data, extension)
            except Exception:
                continue
            if text.strip():
                parts.append(f"\n[{member_path.as_posix()}]\n{text.strip()}")
                extracted_documents += 1
            if (
                extracted_documents >= MAX_CONTAINER_DOCUMENTS
                or len("\n".join(parts)) >= MAX_CONTENT_CHARS
            ):
                break
    return "\n".join(parts)


def _candidate_priority(display_path: str, extension: str) -> tuple[object, ...]:
    """Put reports and shallow readable documents before generated files."""

    normalized = display_path.casefold()
    has_academic_marker = any(
        marker in normalized for marker in ACADEMIC_FILENAME_MARKERS
    )
    is_document = extension in {".docx", ".pdf"}
    is_readable = extension in (PLAIN_TEXT_EXTENSIONS | {".docx", ".pdf"})
    depth = len(PurePosixPath(display_path).parts)
    return (
        not has_academic_marker,
        not is_document,
        not is_readable,
        depth,
        normalized,
    )


def _display_zip_name(name: str) -> str:
    """Best-effort recovery for legacy Chinese ZIP member names."""

    try:
        return name.encode("cp437").decode("gbk")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return name


def _read_document_path(path: Path, extension: str) -> str:
    """Read one supported nested document from disk."""

    if extension in PLAIN_TEXT_EXTENSIONS:
        with path.open("r", encoding="utf-8", errors="replace") as stream:
            return stream.read(MAX_CONTENT_CHARS + 1)
    data = path.read_bytes()
    return _read_document_bytes(data, extension)


def _read_document_bytes(data: bytes, extension: str) -> str:
    """Read DOCX/PDF/text data without extracting an archive to disk."""

    if extension in PLAIN_TEXT_EXTENSIONS:
        return data.decode("utf-8", errors="replace")[: MAX_CONTENT_CHARS + 1]
    if extension == ".docx":
        document = Document(io.BytesIO(data))
        return "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text
        )
    reader = PdfReader(io.BytesIO(data))
    return "\n".join(
        (page.extract_text() or "")
        for page in reader.pages[:MAX_PDF_PAGES]
    )


def _bounded_excerpt(text: str) -> str:
    """Normalize surrounding whitespace and enforce the character limit."""

    return text.strip()[:MAX_CONTENT_CHARS]
