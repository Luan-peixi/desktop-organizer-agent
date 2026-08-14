"""Tests for bounded, local document text extraction."""

from pathlib import Path
import zipfile

from docx import Document
from pypdf import PdfWriter

import desktop_agent.content_extractor as extractor_module
from desktop_agent.content_extractor import extract_content
from desktop_agent.models import ContentExtractionStatus
from desktop_agent.scanner import scan_directory


def _scan_one(directory: Path):
    """Return the only metadata record in a test directory."""

    files = scan_directory(directory)
    assert len(files) == 1
    return files[0]


def test_plain_text_extraction_is_bounded(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Only the configured number of characters should reach metadata."""

    path = tmp_path / "notes.txt"
    path.write_text("abcdefghij")
    monkeypatch.setattr(extractor_module, "MAX_CONTENT_CHARS", 5)

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.EXTRACTED
    assert result.content_excerpt == "abcde"


def test_docx_paragraphs_are_extracted(tmp_path: Path) -> None:
    """Word paragraph text should be available to the classifier."""

    path = tmp_path / "项目计划.docx"
    document = Document()
    document.add_paragraph("火星探测项目里程碑")
    document.add_paragraph("下一阶段完成数据分析")
    document.save(path)

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.EXTRACTED
    assert result.content_excerpt == (
        "火星探测项目里程碑\n下一阶段完成数据分析"
    )


def test_pdf_is_processed_without_exceeding_the_page_limit(
    tmp_path: Path,
) -> None:
    """A valid PDF should be handled locally even when it has no text."""

    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as stream:
        writer.write(stream)

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.EXTRACTED
    assert result.content_excerpt == ""


def test_unsupported_binary_file_is_not_opened(tmp_path: Path) -> None:
    """Installers and other unsupported formats should use metadata only."""

    path = tmp_path / "installer.dmg"
    path.write_bytes(b"not a real installer")

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.UNSUPPORTED
    assert result.content_excerpt is None


def test_broken_document_does_not_stop_the_workflow(tmp_path: Path) -> None:
    """A corrupt supported document should become an error status."""

    path = tmp_path / "broken.docx"
    path.write_text("not a real Word package")

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.ERROR
    assert result.content_excerpt is None


def test_directory_container_extracts_nested_docx_content(
    tmp_path: Path,
) -> None:
    """A neutral folder name should gain evidence from its report body."""

    container = tmp_path / "20260001_示例学生_实验5"
    container.mkdir()
    document = Document()
    document.add_heading("计算机组成原理实验报告", level=1)
    document.add_paragraph("分析 CPU、ALU、寄存器与指令系统。")
    document.save(container / "实验报告.docx")

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.EXTRACTED
    assert "实验报告.docx" in result.content_excerpt
    assert "计算机组成原理" in result.content_excerpt
    assert "CPU" in result.content_excerpt
    assert result.academic_material_hint == "实验报告"


def test_directory_container_reaches_a_report_at_depth_three(
    tmp_path: Path,
) -> None:
    """Two nested experiment folders may contain the actual Word report."""

    container = tmp_path / "20260001_示例学生_实验5"
    report_directory = container / "实验一" / "提交文件"
    report_directory.mkdir(parents=True)
    document = Document()
    document.add_heading("人工智能实验报告", level=1)
    document.add_paragraph("实现智能体状态、工具调用与任务规划。")
    document.save(report_directory / "实验报告.docx")

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.EXTRACTED
    assert "实验一/提交文件/实验报告.docx" in result.content_excerpt
    assert "人工智能实验报告" in result.content_excerpt
    assert result.academic_material_hint == "实验报告"


def test_zip_container_reads_docx_without_extracting_to_disk(
    tmp_path: Path,
) -> None:
    """ZIP reports should be read in memory and leave no extracted files."""

    source_docx = tmp_path / "source.docx"
    document = Document()
    document.add_heading("半导体物理实验", level=1)
    document.add_paragraph("讨论能带、载流子、费米能级与 PN 结。")
    document.save(source_docx)
    archive_path = tmp_path / "20260001_示例学生_实验5.zip"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.write(source_docx, "提交材料/实验报告.docx")
    source_docx.unlink()

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.EXTRACTED
    assert "提交材料/实验报告.docx" in result.content_excerpt
    assert "半导体物理" in result.content_excerpt
    assert list(tmp_path.iterdir()) == [archive_path]


def test_directory_ignores_dependency_tree_and_prioritizes_report(
    tmp_path: Path,
) -> None:
    """A dependency folder must not consume the bounded report budget."""

    container = tmp_path / "20260001_示例学生_实验6"
    dependency = container / "venv" / "lib"
    dependency.mkdir(parents=True)
    for index in range(50):
        (dependency / f"module_{index:02d}.py").write_text("generated code")

    report_directory = container / "提交材料"
    report_directory.mkdir()
    document = Document()
    document.add_heading("人工智能实验报告", level=1)
    document.add_paragraph("实现多智能体协作与工具调用。")
    document.save(report_directory / "实验报告.docx")

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.EXTRACTED
    assert "提交材料/实验报告.docx" in result.content_excerpt
    assert "人工智能实验报告" in result.content_excerpt
    assert "venv/lib" not in result.content_excerpt


def test_large_zip_is_inspected_selectively(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Archive size should not block a small report inside the archive."""

    archive_path = tmp_path / "实验6.zip"
    with zipfile.ZipFile(archive_path, "w", compression=zipfile.ZIP_STORED) as archive:
        archive.writestr("assets/padding.bin", b"x" * 500)
        archive.writestr("提交/实验报告.txt", "人工智能智能体实验")
    assert archive_path.stat().st_size > 100
    monkeypatch.setattr(extractor_module, "MAX_FILE_BYTES", 100)

    result = extract_content(_scan_one(tmp_path))

    assert result.content_status is ContentExtractionStatus.EXTRACTED
    assert "提交/实验报告.txt" in result.content_excerpt
    assert "人工智能智能体实验" in result.content_excerpt
    assert result.academic_material_hint == "实验报告"
