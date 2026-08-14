"""Build a reproducible cross-major demo directory for presentations."""

from __future__ import annotations

import io
import shutil
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIRECTORY = PROJECT_ROOT / "showcase_demo"
FONT_CANDIDATES = (
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
    Path("/System/Library/Fonts/STHeiti Medium.ttc"),
)


def main() -> None:
    """Recreate the presentation-safe sample files."""

    if OUTPUT_DIRECTORY.exists():
        shutil.rmtree(OUTPUT_DIRECTORY)
    OUTPUT_DIRECTORY.mkdir()

    _write_text(
        "课程资料_01.txt",
        "民事诉讼法案例分析\n讨论地域管辖、举证责任、审判程序与执行程序。",
    )
    _write_text(
        "实验记录_02.md",
        "# 半导体物理实验\n\n分析能带、载流子浓度、费米能级与 PN 结伏安特性。",
    )
    _write_text(
        "学习笔记_03.md",
        "# 中国古代文学\n\n本周研读宋词的意象、格律及婉约派与豪放派风格。",
    )
    _write_text(
        "待确认文件",
        "信息不足，故意保留为人工确认样例。",
    )
    _write_docx(
        "小组讨论_04.docx",
        "儿科学病例讨论",
        [
            "病例：儿童发热伴咳嗽三天。",
            "讨论儿童生长发育、鉴别诊断与儿科诊疗方案。",
        ],
    )
    _write_docx(
        "课程作业_05.docx",
        "证券投资学课程作业",
        [
            "比较股票和债券的风险收益特征。",
            "使用投资组合理论分析资产配置方案。",
        ],
    )
    _write_slide_image(
        "课堂截图_06.png",
        "线性代数",
        "矩阵的特征值与特征向量",
        ["A x = λ x", "det(A - λI) = 0", "应用：线性变换与主成分分析"],
        (37, 99, 235),
    )
    _write_slide_image(
        "课堂截图_07.png",
        "集成电路工艺原理",
        "从晶圆到芯片",
        ["光刻  →  刻蚀", "离子注入  →  薄膜生长", "版图、制造、封装与测试"],
        (124, 58, 237),
    )
    _write_slide_image(
        "课堂截图_08.png",
        "人体解剖学",
        "神经系统与主要器官结构",
        ["中枢神经系统", "周围神经系统", "结构、位置与功能关系"],
        (5, 150, 105),
    )
    _write_docx(
        "20260001_示例学生_实验5/实验报告.docx",
        "自动控制原理实验报告",
        [
            "实验内容：建立控制系统传递函数并绘制根轨迹。",
            "使用频率响应和稳定判据分析闭环系统稳定性。",
        ],
    )
    _write_zip_docx(
        "提交材料_10.zip",
        "实验5/实验报告.docx",
        "结构力学实验报告",
        [
            "使用位移法分析超静定结构。",
            "计算杆件内力、节点位移并验证力矩分配结果。",
        ],
    )
    _write_docx(
        "17-智能体/期末大作业/提交版/实验报告.docx",
        "人工智能智能体大作业实验报告",
        [
            "本项目使用 LangGraph 构建多个智能体协作流程。",
            "实现工具调用、状态管理、任务规划与人工审核。",
        ],
    )


def _write_text(name: str, content: str) -> None:
    (OUTPUT_DIRECTORY / name).write_text(content, encoding="utf-8")


def _write_docx(name: str, title: str, paragraphs: list[str]) -> None:
    document = Document()
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    destination = OUTPUT_DIRECTORY / name
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def _write_zip_docx(
    archive_name: str,
    member_name: str,
    title: str,
    paragraphs: list[str],
) -> None:
    """Create a ZIP containing a real Word report without loose temp files."""

    document = Document()
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    with zipfile.ZipFile(
        OUTPUT_DIRECTORY / archive_name,
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as archive:
        archive.writestr(member_name, buffer.getvalue())


def _write_slide_image(
    name: str,
    course: str,
    title: str,
    bullets: list[str],
    accent: tuple[int, int, int],
) -> None:
    image = Image.new("RGB", (1280, 720), "white")
    draw = ImageDraw.Draw(image)
    title_font = _font(56)
    subtitle_font = _font(36)
    body_font = _font(30)
    small_font = _font(22)

    draw.rectangle((0, 0, 1280, 92), fill=accent)
    draw.text((58, 21), course, font=subtitle_font, fill="white")
    draw.text((70, 145), title, font=title_font, fill=(25, 35, 55))
    draw.line((70, 225, 1210, 225), fill=accent, width=5)

    y = 290
    for bullet in bullets:
        draw.ellipse((78, y + 10, 94, y + 26), fill=accent)
        draw.text((120, y), bullet, font=body_font, fill=(45, 55, 75))
        y += 92

    draw.rounded_rectangle(
        (70, 625, 1210, 680),
        radius=14,
        fill=(242, 245, 250),
    )
    draw.text(
        (94, 640),
        "智能桌面整理 Agent · 跨专业视觉识别演示",
        font=small_font,
        fill=(90, 100, 120),
    )
    image.save(OUTPUT_DIRECTORY / name, format="PNG")


def _font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for font_path in FONT_CANDIDATES:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


if __name__ == "__main__":
    main()
